#

import requests
import json
from bs4 import BeautifulSoup
import time
import random
from docx import Document
from docx.shared import Inches
import os

#path = "/data/Cong/es"

## 将文档导入es中进行搜索

def save_to_elasticsearch(id, news_title, news_content):
    headers = {
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6',
        'Content-Type': 'application/json'
    }
    news = {
        "id": str(id),
        "title": str(news_title),
        "content": news_content,
        "category": "高血压"
    }
    print(json.dumps(news))
    search_url = 'http://60.205.159.92:9200/manual/_doc/'+str(id)
    resp = requests.put(
        url=search_url, data=json.dumps(news), headers=headers)
    print(resp.status_code)


def get_case_from_docx():
    files = os.listdir("/data/Cong/manual_docx")
    path= "/data/Cong/manual_docx"
    i=1
    for file in files:
        if not os.path.isdir(file):
            print(file)
            (file_name, extension) = os.path.splitext(file)
            doc = Document(path + "/" + file)
            counter = 0
            news_content = []
            while counter < len(doc.paragraphs):
                news_content.append(doc.paragraphs[counter].text)
                counter += 1
            save_to_elasticsearch(i, file_name, news_content)
            print(str(i)+file_name)
            i=i+1

def get_news():
    path = '/data/Cong/manual_txt'
    files = os.listdir("/data/Cong/manual_txt")
    i = 12
    for file in files:
        if not os.path.isdir(file):
            print(file)
            news_content = []
            with open(path + "/" + file, 'r') as f:
                for line in f.read().splitlines():
                    if len(line) > 0:
                        print(line)
                        news_content.append(str(line))
             # 获取文件名称和后缀名
            (file_name, extension) = os.path.splitext(file)
            save_to_elasticsearch(i,file_name, news_content)
            print(str(i)+"======"+file_name)
            i = i+1
    # 保存到elasticsearch


if __name__ == '__main__':
    get_case_from_docx() # 11
    get_news()
