# 将病历中的数据导入到elasticsearch

import requests
import json
from bs4 import BeautifulSoup
import time
import random
from docx import Document
from docx.shared import Inches
import os

path = "/data/Cong/新增病历/肿瘤科"


def save_to_elasticsearch(news_title, news_date, news_content):
    headers = {
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6',
        'Content-Type': 'application/json'
    }
    news = {
        'title': str(news_title),
        'recordTime': str(news_date),
        'content': news_content,
        'category': '肿瘤科'
    }
    print(news)
    search_url = 'http://127.0.0.1:15002/cases/save'
    resp = requests.post(url=search_url, data=json.dumps(news), headers=headers)
    print(resp.status_code)


def get_news():
    files = os.listdir(path)
    for file in files:
        if not os.path.isdir(file):
            print(file)
            doc = Document(path + "/" + file)
            counter = 0
            news_title = ''
            news_date = ''
            news_content = []
            while counter < len(doc.paragraphs):
                if counter == 0:
                    news_title = doc.paragraphs[counter].text
                elif counter == 1:
                    news_date = doc.paragraphs[counter].text
                elif len(doc.paragraphs[counter].text) > 0:
                    news_content.append(doc.paragraphs[counter].text)
                counter += 1
            save_to_elasticsearch(news_title, news_date, news_content)
            print(news_title + news_date)
    # 保存到elasticsearch
    # 创建word文档


if __name__ == '__main__':
    get_news()
