# 百度新闻接口数据爬取

import requests
import json
from bs4 import BeautifulSoup
import time
import random
from docx import Document
from docx.shared import Inches
import os

path = "/data/Space/guid"


def save_to_elasticsearch(guid_title, guid_date, guid_content):
    headers = {
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6',
        'Content-Type': 'application/json'
    }
    guid = {
        'title': str(guid_title),
        'guidDate': str(guid_date),
        'content': str(guid_content),
        'category': '国内疫情'
    }
    print(guid)
    search_url = 'http://127.0.0.1:15002/covid/guid/save'
    resp = requests.post(url=search_url, data=json.dumps(guid), headers=headers)
    print(resp.status_code)


def get_guid():
    files = os.listdir(path)
    for file in files:
        if not os.path.isdir(file):
            print(file)
            doc = Document(path + "/" + file)
            counter = 0
            while counter < len(doc.paragraphs):
                print(doc.paragraphs[counter].text)
                print(counter)
                counter += 1
            # for p in doc.paragraphs:
            #     print(len(p))
            #     print(p.text)

            # save_to_elasticsearch(guid_title,guid_date,guid_content)

    # 保存到elasticsearch
    # 创建word文档


if __name__ == '__main__':
    get_guid()
