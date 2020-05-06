from docx import Document
from docx.shared import Inches
doc = Document()
doc.add_heading("这是一个标题")
doc.add_paragraph("段落")
doc.add_picture('./33.png',width=Inches(5.0))
doc.save('demo.docx')
