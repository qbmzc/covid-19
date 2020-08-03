# 百度新闻接口数据爬取

import requests
import json
from bs4 import BeautifulSoup
import time
import random
from docx import Document
from docx.shared import Inches


def get_news(req_url, eventDescription, eventTime):
    headers = {
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6',
        'Content-Type': 'application/json'
    }
    doc = Document('/data/Space/covid/baidu/2020-07-11 15:24:52_国家卫健委：本次聚集性疫情新增本土确诊病例占全国新增确诊病例68%.docx')
    for d in doc.paragraph:
        print(d)
    doc.add_heading()
    doc.add_paragraph(news_date)
    doc.add_paragraph(news_content)
    for img in data.find_all('img', class_='large'):
        ssrc = img.get('src')

        print(ssrc)

        download_img(ssrc)
        doc.add_picture('./img.jpg', width=Inches(5.0), height=Inches(5.0))
    doc.save(directory + news_date + "_" + news_title.replace("/", "_") + '.docx')  # 标题中 不能存在`/`
    # 保存到elasticsearch
    # 创建word文档
    news = {
        'title': str(news_title),
        'newDate': str(news_date),
        'content': str(news_content),
        'category': 'news'
    }
    print(news)
    search_url = 'http://127.0.0.1:15002/covid/news/save'
    resp = requests.post(url=search_url, data=json.dumps(news), headers=headers)
    print(resp.status_code)


if __name__ == '__main__':
        get_news(eventUrl, eventDescription, eventTime)
