# 获取中国疾控中心新闻数据
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches
import time
import random

base_url = 'http://www.chinacdc.cn/'
gwxx_url = base_url + 'yw_9324'
base_directory = '/data/Space/covid/cdc/yw_9324/'
# 页面
pape_count = 13


def get_data(url):
    headers = {
        'Host': 'www.chinacdc.cn',
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6'
    }
    r = requests.get(url, headers=headers)
    # print(r.status_code)
    if r.status_code == 200:
        soup = BeautifulSoup(r.text, "html.parser")
        # print(soup)
        return soup


# 定义一个筛选器
def has_span_in_p(tag):
    return tag.has_attr('style')


def get_news(req_url):
    data = get_data(req_url)
    a_list = data.find_all('a', class_='item-text')
    for a in a_list:
        a_href = a.get('href')
        # 拼接正常的路径
        s_href = a_href.replace('.', gwxx_url, 1)
        news_soup = get_data(s_href)
        news_title = news_soup.find("p", class_='cn-main-title').text
        title = news_title.replace('/', '_')
        news_date = news_soup.find('span', class_='info-date').text
        news_content = ''
        for cc in news_soup.find_all('span'):
            news_content += cc.text

        # 创建word文档
        doc = Document()
        doc.add_heading(news_title)
        doc.add_paragraph(news_date)
        doc.add_paragraph(news_content)
        for img in news_soup.find_all('img'):
            ssrc = img.get('src')
            if ssrc.find('images') == -1:
                print(ssrc)
                real_path_0 = img.get('src').replace('.', '', 1)
                center_path = a_href.split('/')[1]
                print(center_path)
                real_path = gwxx_url + "/" + center_path + real_path_0
                print(real_path)
                download_img(real_path)
                doc.add_picture('/data/Space/covid-19/src/img.png', width=Inches(5.0), height=Inches(5.0))
        doc.save(base_directory + news_date + "_" + title + '.docx')


# http://www.chinacdc.cn/gwxx/index_12.html


def download_img(image_url):
    headers = {
        'Host': 'www.chinacdc.cn',
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6'
    }
    r = requests.get(image_url, headers=headers)
    if r.status_code == 200:
        with open('./img.png', 'wb') as f:
            f.write(r.content)
        print('done')
        time.sleep(random.uniform(0, 2))
    del r


if __name__ == '__main__':
    for i in range(7, 13):
        url = gwxx_url + '/index_' + str(i) + ".html"
        print(url)
        time.sleep(2)
        get_news(url)
