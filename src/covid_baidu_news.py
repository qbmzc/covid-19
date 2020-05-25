# 百度新闻接口数据爬取

import requests
import json
from bs4 import BeautifulSoup
import time
import random
from docx import Document
from docx.shared import Inches

# 肺炎
# 新冠肺炎国外疫情

def getData():
    url = "https://opendata.baidu.com/data/inner?tn=reserved_all_res_tn&dspName=iphone&from_sf=1&dsp=iphone&resource_id=28565&alr=1&query=肺炎"
    headers = {
        'user-agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.113 '
                      'Safari/537.36 '
    }
    r = requests.get(url, headers=headers)
    if r.status_code == 200:
        return json.loads(r.text)['Result']


data_news = getData()


def getEventUrl(data_news):
    for i in data_news:
        ss = i['items_v2']
        for ii in ss:
            s = ii['aladdin_res']['DisplayData']['result']
            r = s['items']
            return r


def get_data(url):
    headers = {
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6'
    }
    r = requests.get(url, headers=headers)
    # print(r.status_code)
    if r.status_code == 200:
        soup = BeautifulSoup(r.text, "html.parser")
        # print(soup)
        return soup


# print(data_news['result'])

# print(data_news)

# lastUpdateTime = str(datetime.date.today())  #
directory = "/data/Space/covid/baidu/"  # 定义数据保存路径

# 图片下载
def download_img(image_url):
    headers = {
        'Connection': 'keep-alive',
        'Pragma': 'no-cache',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.122 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Accept-Language': 'zh-CN,zh;q=0.9,en-US;q=0.8,en;q=0.7,zh-TW;q=0.6'
    }
    r = requests.get(image_url, headers=headers)
    if r.status_code == 200:
        with open('./img.jpg', 'wb') as f:
            f.write(r.content)
        print('done')
        time.sleep(random.uniform(0, 2))
    del r


def get_news(req_url, eventDescription, eventTime):
    data = get_data(req_url)
    a_list = data.find_all('span', class_='bjh-p')
    news_content = ''
    for a in a_list:
        news_content += a.text
        # 拼接正常的路径
    news_title = eventDescription
    time_local = time.localtime(int(eventTime))
    news_date = time.strftime("%Y-%m-%d %H:%M:%S", time_local)
    # 创建word文档
    doc = Document()
    doc.add_heading(news_title)
    doc.add_paragraph(news_date)
    doc.add_paragraph(news_content)
    for img in data.find_all('img', class_='large'):
        ssrc = img.get('src')

        print(ssrc)

        download_img(ssrc)
        doc.add_picture('./img.jpg', width=Inches(5.0), height=Inches(5.0))
    doc.save(directory + news_date + "_" + news_title + '.docx')


if __name__ == '__main__':
    url_list = getEventUrl(getData())
    for urls in url_list:
        eventUrl = urls['eventUrl']
        eventDescription = urls['eventDescription']
        eventTime = urls['eventTime']
        print(eventUrl+eventDescription+eventTime)
        get_news(eventUrl, eventDescription, eventTime)
